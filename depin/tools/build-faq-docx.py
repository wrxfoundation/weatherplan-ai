#!/usr/bin/env python3
"""wellbian FAQ → DOCX.

lib/data.ts 를 직접 파싱해 만들므로 사이트 원본과 어긋나지 않는다.
1차본은 전부 Normal 스타일 + 직접서식이라 탐색창·목차·개요가 잡히지 않았다 →
Heading 스타일 / 목차 필드 / 머리글·바닥글(쪽번호) 을 갖춘 문서로 재작성.
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm

SRC = Path(sys.argv[1] if len(sys.argv) > 1 else "lib/data.ts")
OUT = Path(sys.argv[2] if len(sys.argv) > 2 else "public/wellbian-faq.docx")
ASOF = "2026-08-28"

KO_FONT = "Malgun Gothic"
INK = RGBColor(0x1B, 0x1B, 0x48)
MUTE = RGBColor(0x6B, 0x6B, 0x8C)
BRAND = RGBColor(0x4F, 0x46, 0xE5)

# ── data.ts 파싱 ────────────────────────────────────────────────────────────
ITEM = re.compile(r'\{\s*q:\s*"((?:[^"\\]|\\.)*)"\s*,\s*a:\s*"((?:[^"\\]|\\.)*)"\s*\}')


def unesc(s: str) -> str:
    return s.replace('\\"', '"').replace("\\\\", "\\").replace("\\n", " ")


def grab(src: str, name: str):
    m = re.search(r"export const %s\s*(?::[^=]*)?=\s*\[" % re.escape(name), src)
    if not m:
        raise SystemExit(f"배열을 찾지 못함: {name}")
    i = m.end() - 1
    depth, j = 0, i
    while j < len(src):
        if src[j] == "[":
            depth += 1
        elif src[j] == "]":
            depth -= 1
            if depth == 0:
                break
        j += 1
    body = src[i : j + 1]
    return [(unesc(q), unesc(a)) for q, a in ITEM.findall(body)]


src = SRC.read_text(encoding="utf-8")
KO_CORE = grab(src, "FAQS")
KO_MORE = grab(src, "FAQS_EXTRA")
EN_CORE = grab(src, "FAQS_EN")
EN_MORE = grab(src, "FAQS_EXTRA_EN")

assert len(KO_CORE) == len(EN_CORE), f"기본 문항 수 불일치 KO {len(KO_CORE)} / EN {len(EN_CORE)}"
assert len(KO_MORE) == len(EN_MORE), f"확장 문항 수 불일치 KO {len(KO_MORE)} / EN {len(EN_MORE)}"
TOTAL = len(KO_CORE) + len(KO_MORE)

# ── 문서 골격 ──────────────────────────────────────────────────────────────
doc = Document()


def set_font(style, name=KO_FONT, size=None, bold=None, color=None):
    """한글은 rFonts의 eastAsia까지 지정해야 Word에서 대체 글꼴로 튀지 않는다."""
    style.font.name = name
    if size is not None:
        style.font.size = Pt(size)
    if bold is not None:
        style.font.bold = bold
    if color is not None:
        style.font.color.rgb = color
    rpr = style.element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rpr.append(rf)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rf.set(qn(attr), name)


s = doc.styles
set_font(s["Normal"], size=10.5, color=INK)
s["Normal"].paragraph_format.space_after = Pt(6)
s["Normal"].paragraph_format.line_spacing = 1.32

set_font(s["Title"], size=24, bold=True, color=INK)
set_font(s["Heading 1"], size=16, bold=True, color=BRAND)
set_font(s["Heading 2"], size=13, bold=True, color=INK)
set_font(s["Heading 3"], size=11, bold=True, color=INK)
for h, before, after in (("Heading 1", 22, 8), ("Heading 2", 16, 6), ("Heading 3", 12, 3)):
    pf = s[h].paragraph_format
    pf.space_before, pf.space_after = Pt(before), Pt(after)
    pf.keep_with_next = True

sec = doc.sections[0]
sec.left_margin = sec.right_margin = Cm(2.4)
sec.top_margin = Cm(2.2)
sec.bottom_margin = Cm(2.0)


def field(par, instr: str, placeholder: str = ""):
    """Word 필드(목차·쪽번호)를 raw OXML로 삽입."""
    r = par.add_run()
    fc = OxmlElement("w:fldChar")
    fc.set(qn("w:fldCharType"), "begin")
    r._r.append(fc)

    r = par.add_run()
    it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve")
    it.text = instr
    r._r.append(it)

    r = par.add_run()
    fc = OxmlElement("w:fldChar")
    fc.set(qn("w:fldCharType"), "separate")
    r._r.append(fc)

    if placeholder:
        par.add_run(placeholder)

    r = par.add_run()
    fc = OxmlElement("w:fldChar")
    fc.set(qn("w:fldCharType"), "end")
    r._r.append(fc)


# 머리글 / 바닥글
hp = sec.header.paragraphs[0]
hp.text = "wellbian ARC-600DA — 자주 묻는 질문(FAQ)"
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
for r in hp.runs:
    r.font.size, r.font.color.rgb, r.font.name = Pt(8.5), MUTE, KO_FONT

fp = sec.footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
field(fp, " PAGE ", "1")
for r in fp.runs:
    r.font.size, r.font.color.rgb, r.font.name = Pt(9), MUTE, KO_FONT

# ── 표지 블록 ──────────────────────────────────────────────────────────────
doc.add_paragraph("wellbian ARC-600DA — 자주 묻는 질문(FAQ)", style="Title")

sub = doc.add_paragraph()
r = sub.add_run("날씨데이터토큰생성기™ · 사전예약 안내 문서")
r.font.size, r.font.color.rgb = Pt(12), MUTE

meta = doc.add_paragraph()
for line in (
    f"기준일 {ASOF}",
    "판매 Wellbian Labs Pte. Ltd.(싱가포르)",
    "기기 파트너 · 국내 수탁 케이웨더",
    f"총 {TOTAL}문항 (기본 {len(KO_CORE)} · 전체 확장 {len(KO_MORE)}) · 국문/영문 병기",
):
    rr = meta.add_run(line + "\n")
    rr.font.size, rr.font.color.rgb = Pt(9.5), MUTE

note = doc.add_paragraph()
rn = note.add_run(
    "본 문서는 사이트(weatherplan-ai / wellbian store)의 FAQ 원본 데이터에서 자동 생성됩니다. "
    "문구를 수정할 때는 이 문서가 아니라 사이트 원본을 고쳐 주세요."
)
rn.font.size, rn.font.color.rgb, rn.italic = Pt(9), MUTE, True

# ── 목차 ───────────────────────────────────────────────────────────────────
doc.add_paragraph("목차", style="Heading 1")
toc = doc.add_paragraph()
field(toc, r' TOC \o "1-2" \h \z \u ', "[Word에서 F9를 누르면 목차가 채워집니다]")
for r in toc.runs:
    if r.text.startswith("["):
        r.font.size, r.font.color.rgb, r.italic = Pt(9), MUTE, True

doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def qa_block(items, prefix_from=1):
    for n, (q, a) in enumerate(items, prefix_from):
        doc.add_heading(f"Q{n}. {q}", level=3)
        p = doc.add_paragraph(a)
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(10)


def caption(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size, r.font.color.rgb = Pt(9), MUTE
    p.paragraph_format.space_after = Pt(10)


# ── 국문 ───────────────────────────────────────────────────────────────────
doc.add_heading("국문 FAQ", level=1)

doc.add_heading(f"기본 FAQ ({len(KO_CORE)}문항)", level=2)
caption("사이트 FAQ 섹션에 기본 노출되는 문항입니다.")
qa_block(KO_CORE, 1)

doc.add_heading(f"전체 FAQ · 확장 ({len(KO_MORE)}문항)", level=2)
caption("'전체 FAQ 보기'를 눌렀을 때 펼쳐지는 문항입니다.")
qa_block(KO_MORE, len(KO_CORE) + 1)

doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

# ── 영문 ───────────────────────────────────────────────────────────────────
doc.add_heading("English FAQ", level=1)

doc.add_heading(f"Core FAQ ({len(EN_CORE)} questions)", level=2)
caption("Shown by default in the site's FAQ section.")
qa_block(EN_CORE, 1)

doc.add_heading(f"Full FAQ · expanded ({len(EN_MORE)} questions)", level=2)
caption("Revealed by the 'See all FAQs' toggle.")
qa_block(EN_MORE, len(EN_CORE) + 1)

# python-docx 기본 템플릿의 settings.xml 은 <w:zoom w:val="bestFit"/> 만 넣는데
# 스키마상 w:percent 가 필수라 검증에서 걸린다. Word는 눈감아 주지만 채워 둔다.
zoom = doc.settings.element.find(qn("w:zoom"))
if zoom is not None and zoom.get(qn("w:percent")) is None:
    zoom.set(qn("w:percent"), "100")

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(f"저장: {OUT} ({OUT.stat().st_size:,}B)")
print(f"국문 {len(KO_CORE)}+{len(KO_MORE)} · 영문 {len(EN_CORE)}+{len(EN_MORE)} = 총 {TOTAL}문항 (KO/EN 각각)")
